#!/usr/bin/env python3
"""
Create a reference DOCX document with custom styles for Pandoc
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.shared import OxmlElement, qn

def create_reference_docx():
    doc = Document()
    
    # Set default font to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    
    # Set double spacing
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Create Heading 1 style (11pt, Bold)
    heading1 = doc.styles['Heading 1']
    heading1_font = heading1.font
    heading1_font.name = 'Times New Roman'
    heading1_font.size = Pt(11)
    heading1_font.bold = True
    
    # Create Heading 2 style (10pt)
    heading2 = doc.styles['Heading 2']
    heading2_font = heading2.font
    heading2_font.name = 'Times New Roman'
    heading2_font.size = Pt(10)
    
    # Create Heading 3 style (10pt)
    heading3 = doc.styles['Heading 3']
    heading3_font = heading3.font
    heading3_font.name = 'Times New Roman'
    heading3_font.size = Pt(10)
    
    # Add some sample content to establish styles
    doc.add_heading('Sample Heading 1', level=1)
    doc.add_heading('Sample Heading 2', level=2)
    doc.add_heading('Sample Heading 3', level=3)
    doc.add_paragraph('This is sample normal text in Times New Roman 9pt with double spacing.')
    
    # Save the reference document
    doc.save('reference.docx')
    print("Reference document created: reference.docx")

if __name__ == "__main__":
    create_reference_docx() 